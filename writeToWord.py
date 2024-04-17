from docx import Document
from docx.shared import Inches

def create_word_document(pictures, output_file):
    document = Document()

    for picture in pictures:
        # Add picture to paragraph
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(picture, width=Inches(3))

        # Center-align paragraph
        paragraph.alignment = 1  # Set alignment to center (0=left, 1=center, 2=right)

    # Save document
    document.save(output_file)
    print(f"Word document '{output_file}' created successfully.")

# Example usage
pictures = ["pic1.jpg", "pic2.jpg", "pic3.png"]
output_file = "output_document.docx"
create_word_document(pictures, output_file)
